#!/usr/bin/env python
"""Build Supplementary Table Mouse 1 — multi-age single-cell reference assembly
for the whole-body spatial AD-risk analysis.

Outputs (in the same directory as this script):
  supplement_table_mouse_1.md                       — narrative markdown
  supplement_table_mouse_1_sources.csv              — Section 1: 5 source datasets
  supplement_table_mouse_1_summary.csv              — Section 2: per-source contribution
  supplement_table_mouse_1_ledger.csv               — Section 3: full per-(organ × age × source) table
  supplement_table_mouse_1_notes.csv                — Notes
  supplement_table_mouse_1.xlsx                     — multi-sheet workbook
  supplement_table_mouse_1.docx                     — formatted Word document

Driving data: results/age_sample_provenance_per_source.csv (built by
scripts/build_age_provenance_per_source.py).

Run: python tables/build_supp_table_mouse_1.py
"""
import os
from pathlib import Path

import pandas as pd

# ------------------------------------------------------------
# Paths
# ------------------------------------------------------------
REPO = Path(__file__).resolve().parents[1]
OUT  = REPO / "tables"
PROV = REPO / "results" / "age_sample_provenance_per_source.csv"

OUT.mkdir(parents=True, exist_ok=True)

# ------------------------------------------------------------
# Static metadata
# ------------------------------------------------------------
SRC_META = {
    "PanSci": dict(
        paper="Zhang et al. Nature 2024",
        accession="GSE247719",
        assay="sci-RNA-seq3 (PanSci)",
        organs="Kidney, Lung, Heart, Liver, Muscle, Stomach, BAT, iWAT, gWAT, "
               "Ileum, Colon, Pancreas, Skin, Thymus",
        ages="3, 6, 12, 16, 23 months",
        preprocess="Curated cellType + Lineage + Sub_cell_type annotations from source",
    ),
    "EasySci": dict(
        paper="Sziraki et al. Nature Aging 2023",
        accession="GSE212606",
        assay="EasySci (sci-RNA-seq3 variant, brain-optimised)",
        organs="Brain (whole)",
        ages="3, 6, 21 months",
        preprocess="Main_cell_type + Sub_cluster from source",
    ),
    "TMS": dict(
        paper="Schaum et al. Nature 2020 (Tabula Muris Senis)",
        accession="figshare:8273102 (TMS)",
        assay="Smart-seq2 (FACS) + 10x Chromium (Droplet)",
        organs="23 organs (BAT, Bone Marrow, Brain Myeloid/NonMyeloid, Heart, "
               "Kidney, Liver, Lung, Pancreas, Spleen, Skin, Muscle, Stomach, "
               "Colon, Thymus, Trachea, …)",
        ages="1, 3, 18, 21, 24, 30 months",
        preprocess="cell_ontology_class from source (Tabula Muris cell ontology)",
    ),
    "MCA2": dict(
        paper="Han et al. Cell 2022 (Mouse Cell Atlas 2.0)",
        accession="GSE153562",
        assay="Microwell-seq (MCA2)",
        organs="Stomach, Ileum, Spleen, Thymus (mid-late life gaps for TMS organs)",
        ages="6, 12, 18, 21, 24 months (per-organ subset)",
        preprocess="leiden clusters with marker-based manual annotation",
    ),
    "Hammond_microglia": dict(
        paper="Olah et al. Nature Communications 2020",
        accession="GSE179358",
        assay="10x Chromium (microglia FACS-sorted, brain)",
        organs="Brain (microglia only)",
        ages="6, 12, 18 months",
        preprocess="cellType=microglial cell (single-cell-type focus)",
    ),
}

ST_TARGET = {
    "BAT":         "Brown_Fat_CTRL1 (GSE248904)",
    "Bone_Marrow": "Bone_Marrow_CTRL1 (GSE248904)",
    "Brain":       "Brain_CTRL1 (GSE248904)",
    "Colon":       "Colon_CTRL1 (GSE248904)",
    "Heart":       "Heart_CTRL1 (GSE248904)",
    "Ileum":       "Small_Intestine_CTRL1 (GSE248904)",
    "Kidney":      "Kidney_CTRL1 (GSE248904)",
    "Liver":       "Liver_CTRL1 (GSE248904)",
    "Lung":        "Lung_CTRL1 (GSE248904)",
    "Muscle":      "Muscle_CTRL1 (GSE248904)",
    "Pancreas":    "Pancreas_CTRL1 (GSE248904)",
    "Skin":        "Skin_CTRL1 (GSE248904)",
    "Spleen":      "Spleen_CTRL1 (GSE248904)",
    "Stomach":     "Stomach_CTRL1 (GSE248904)",
    "Thymus":      "Thymus_CTRL1 (GSE248904)",
}

NOTES = [
    ("CMap projection",
     "Each single-cell sample is mapped onto the whole-body spatial reference "
     "(GSE248904 CTRL replicate 1) using CMap classification + spot assignment. "
     "See scripts/preprocess_new_sc.py and scripts/process_remaining_gaps.py. "
     "A spot-classification confidence (CMAP_quality) is retained per cell."),
    ("gsMap input",
     "Each (organ × age × source) projected h5ad is one input to "
     "`gsmap quick_mode`; outputs include per-spot AD -log10(p) and "
     "per-annotation Cauchy combination."),
    ("Multi-source Brain samples",
     "5 Brain samples (Brain_3m, Brain_6m, Brain_12m, Brain_18m, Brain_24m) "
     "each combine ≥2 sources to maximise cell-type coverage: PanSci/EasySci "
     "supply non-microglia neurons; Hammond and TMS add microglia depth; MCA2 "
     "fills aged gaps. See Fig. 4c pie-matrix in notebooks/paper_figure.ipynb."),
    ("Why 15 organs and not 16",
     "Lymph_Node spatial spots exist in GSE248904 but no aged single-cell "
     "dataset covers lymph node directly; lymph node was excluded from the "
     "multi-age atlas (still appears in the 6-week whole-body baseline)."),
    ("Why 1–30 mo and not 1–32 mo",
     "TMS droplet covers up to 30 m; no public single-cell mouse atlas "
     "reliably covers 31–32 months for the organs analysed here."),
    ("Provenance script",
     "scripts/build_age_provenance_per_source.py rebuilds "
     "results/age_sample_provenance_per_source.csv from raw obs metadata + "
     "CMap output paths."),
]

# ------------------------------------------------------------
# Build dataframes
# ------------------------------------------------------------
df = pd.read_csv(PROV)
df["source_paper"]      = df["source"].map(lambda s: SRC_META.get(s, {}).get("paper", "?"))
df["source_accession"]  = df["source"].map(lambda s: SRC_META.get(s, {}).get("accession", "?"))
df["preprocess_status"] = df["source"].map(lambda s: SRC_META.get(s, {}).get("preprocess", "?"))
df["mapping_target"]    = df["organ"].map(ST_TARGET).fillna("?")
df = df.sort_values(["organ", "age_m", "source"]).reset_index(drop=True)

# Section 1 — sources
df_sources = pd.DataFrame([
    dict(rank=i,
         source=k,
         paper=v["paper"],
         accession=v["accession"],
         assay=v["assay"],
         organs_covered=v["organs"],
         ages_offered=v["ages"],
         preprocessing_status=v["preprocess"])
    for i, (k, v) in enumerate(SRC_META.items(), 1)
])

# Section 2 — per-source contribution
df_summary = (df.groupby("source")
                .agg(samples=("sample", "nunique"),
                     distinct_organs=("organ", "nunique"),
                     total_cells=("n_cells", "sum"),
                     age_range=("age_m", lambda s: f"{int(s.min())}–{int(s.max())} mo"))
                .reset_index()
                .sort_values("total_cells", ascending=False))
df_summary["pct_of_atlas"] = (100 * df_summary["total_cells"]
                              / df_summary["total_cells"].sum()).round(2)
# Append total row
total_row = pd.DataFrame([dict(
    source="Total",
    samples=df["sample"].nunique(),
    distinct_organs=df["organ"].nunique(),
    total_cells=int(df["n_cells"].sum()),
    age_range=f"{int(df['age_m'].min())}–{int(df['age_m'].max())} mo",
    pct_of_atlas=100.0,
)])
df_summary = pd.concat([df_summary, total_row], ignore_index=True)

# Section 3 — full ledger
df_ledger = df[["organ", "age_m", "source", "sample", "n_cells",
                "source_accession", "preprocess_status", "mapping_target"]].copy()
df_ledger = df_ledger.rename(columns={"age_m": "age_months", "sample": "sample_id"})

# Notes
df_notes = pd.DataFrame(NOTES, columns=["topic", "note"])

# ------------------------------------------------------------
# Write CSVs
# ------------------------------------------------------------
df_sources.to_csv(OUT / "supplement_table_mouse_1_sources.csv", index=False)
df_summary.to_csv(OUT / "supplement_table_mouse_1_summary.csv", index=False)
df_ledger.to_csv(OUT / "supplement_table_mouse_1_ledger.csv", index=False)
df_notes.to_csv(OUT / "supplement_table_mouse_1_notes.csv", index=False)

# ------------------------------------------------------------
# XLSX (multi-sheet, formatted)
# ------------------------------------------------------------
xlsx_path = OUT / "supplement_table_mouse_1.xlsx"
with pd.ExcelWriter(xlsx_path, engine="openpyxl") as xw:
    df_sources.to_excel(xw, sheet_name="1 - Source datasets", index=False)
    df_summary.to_excel(xw, sheet_name="2 - Per-source contribution", index=False)
    df_ledger.to_excel(xw, sheet_name="3 - Full ledger", index=False)
    df_notes.to_excel(xw, sheet_name="Notes", index=False)

# Apply lightweight formatting
from openpyxl import load_workbook
from openpyxl.styles import Font, PatternFill, Alignment
wb = load_workbook(xlsx_path)
HEADER_FILL = PatternFill("solid", fgColor="1F4E79")
HEADER_FONT = Font(bold=True, color="FFFFFF", size=10)
BODY_FONT   = Font(size=10)
ALIGN       = Alignment(wrap_text=True, vertical="top")
for sheet in wb.sheetnames:
    ws = wb[sheet]
    for cell in ws[1]:
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = Alignment(wrap_text=True, vertical="center")
    ws.freeze_panes = "A2"
    # Auto-ish width: cap at 50
    for col_cells in ws.columns:
        m = max((len(str(c.value)) if c.value else 0) for c in col_cells)
        ws.column_dimensions[col_cells[0].column_letter].width = min(max(m + 2, 12), 50)
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.font = BODY_FONT
            cell.alignment = ALIGN
wb.save(xlsx_path)

# ------------------------------------------------------------
# DOCX
# ------------------------------------------------------------
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def add_df_table(doc, df_in, header_hex="1F4E79", body_hex=None,
                 cap_widths=None):
    table = doc.add_table(rows=1 + len(df_in), cols=len(df_in.columns))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, col in enumerate(df_in.columns):
        hdr[i].text = str(col)
        for run in hdr[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)
        shade_cell(hdr[i], header_hex)
    for r, (_, row) in enumerate(df_in.iterrows(), 1):
        for i, val in enumerate(row.values):
            table.rows[r].cells[i].text = "" if pd.isna(val) else str(val)
            for run in table.rows[r].cells[i].paragraphs[0].runs:
                run.font.size = Pt(8.5)
    if cap_widths:
        for i, w in enumerate(cap_widths):
            if w is None: continue
            for cell in table.columns[i].cells:
                cell.width = Cm(w)
    return table


doc = Document()
# Landscape A4 for readability of the wide ledger
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width, section.page_height = section.page_height, section.page_width
section.top_margin = section.bottom_margin = Cm(1.2)
section.left_margin = section.right_margin = Cm(1.2)

# Default font
for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3"):
    if style_name in doc.styles:
        font = doc.styles[style_name].font
        font.name = "Arial"

doc.add_heading("Supplementary Table Mouse 1", level=1)
p = doc.add_paragraph()
p.add_run("Source datasets and per-(organ × age) reference assembly for the "
          "whole-body spatial AD-risk analysis. "
          "809,284 cells across 126 organ–age combinations, projected onto the "
          "6-week whole-body Array-seq spatial reference (GSE248904) via CMap. "
          "Each sample feeds gsMap spatial-LDSC to produce per-spot AD heritability."
          ).font.size = Pt(10)

doc.add_heading("1 | Source datasets (5)", level=2)
add_df_table(doc, df_sources, cap_widths=[1.2, 2.6, 4.5, 3.0, 4.5, 6.5, 4.5, 5.0])

doc.add_heading("2 | Per-source contribution", level=2)
add_df_table(doc, df_summary)

doc.add_heading("3 | Full per-(organ × age × source) ledger", level=2)
p = doc.add_paragraph()
p.add_run("131 rows; 5 are multi-source Brain samples where ≥1 dataset is "
          "co-mapped to the same spatial reference.").font.italic = True
add_df_table(doc, df_ledger,
             cap_widths=[2.2, 1.8, 2.5, 4.0, 1.8, 4.0, 6.0, 5.5])

doc.add_heading("Notes", level=2)
for topic, note in NOTES:
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(f"{topic}. ")
    r1.font.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(note)
    r2.font.size = Pt(10)

doc.add_heading("Summary statistics", level=2)
stats = [
    f"Total single-cell sources: {len(SRC_META)}",
    f"Total samples (organ × age × source rows): {len(df_ledger)}",
    f"Distinct (organ × age) combinations: {df['sample'].nunique()}",
    f"Distinct organs: {df['organ'].nunique()}",
    f"Distinct ages: {df['age_m'].nunique()} "
    f"(1, 3, 6, 12, 16, 18, 21, 23, 24, 30 months)",
    f"Total cells: {int(df['n_cells'].sum()):,}",
    f"Largest source: PanSci "
    f"({int(df.loc[df['source']=='PanSci','n_cells'].sum()):,} cells, "
    f"{100*df.loc[df['source']=='PanSci','n_cells'].sum()/df['n_cells'].sum():.1f}%)",
    f"Smallest source: Hammond_microglia "
    f"({int(df.loc[df['source']=='Hammond_microglia','n_cells'].sum()):,} cells)",
    f"Multi-source samples: "
    f"{(df.groupby('sample')['source'].nunique() > 1).sum()} (all Brain)",
]
for s in stats:
    doc.add_paragraph(s, style="List Bullet")

doc.save(OUT / "supplement_table_mouse_1.docx")

# ------------------------------------------------------------
# Markdown (master narrative)
# ------------------------------------------------------------
def fmt_n(n):
    if isinstance(n, float) and n.is_integer(): n = int(n)
    if isinstance(n, int) and n >= 1000: return f"{n:,}"
    return str(n)

with open(OUT / "supplement_table_mouse_1.md", "w") as f:
    f.write("## Supplementary Table Mouse 1 | Source datasets and per-(organ × age) "
            "reference assembly\n\n")
    f.write("Multi-age single-cell reference atlas assembled to provide cell-type and age "
            "context for the whole-body spatial AD-risk analysis. **809,284 cells** across "
            "**126 organ–age combinations**, projected onto the 6-week whole-body Array-seq "
            "spatial reference (GSE248904) using CMap. Each organ–age sample feeds gsMap "
            "spatial-LDSC to produce per-spot AD heritability enrichment.\n\n")

    f.write("### 1 | Source datasets (5)\n\n")
    f.write("| # | Source | Paper | Accession | Assay | Organs covered | Ages offered | "
            "Preprocessing status |\n")
    f.write("|---|---|---|---|---|---|---|---|\n")
    for _, r in df_sources.iterrows():
        f.write(f"| {r['rank']} | **{r['source']}** | {r['paper']} | "
                f"`{r['accession']}` | {r['assay']} | {r['organs_covered']} | "
                f"{r['ages_offered']} | {r['preprocessing_status']} |\n")

    f.write("\n### 2 | Per-source contribution to the assembled reference\n\n")
    f.write("| Source | Samples | Distinct organs | Total cells | % of atlas | Age range |\n")
    f.write("|---|---:|---:|---:|---:|---|\n")
    for _, r in df_summary.iterrows():
        prefix = "**" if r["source"] == "Total" else ""
        suffix = "**" if r["source"] == "Total" else ""
        f.write(f"| {prefix}{r['source']}{suffix} | {prefix}{r['samples']}{suffix} | "
                f"{prefix}{r['distinct_organs']}{suffix} | "
                f"{prefix}{fmt_n(r['total_cells'])}{suffix} | "
                f"{prefix}{r['pct_of_atlas']:.1f}%{suffix} | "
                f"{prefix}{r['age_range']}{suffix} |\n")

    f.write("\n### 3 | Full per-(organ × age × source) ledger (131 rows)\n\n")
    f.write("Multi-source rows appear where >1 dataset contributes cells to the same "
            "organ × age combination (5 such Brain samples). Each row corresponds to one "
            "CMap-projection unit feeding gsMap.\n\n")
    f.write("| Organ | Age (mo) | Source | Sample ID | n_cells | Source accession | "
            "Preprocessing | Mapping target |\n")
    f.write("|---|---:|---|---|---:|---|---|---|\n")
    for _, r in df_ledger.iterrows():
        f.write(f"| {r['organ']} | {int(r['age_months'])} | {r['source']} | "
                f"`{r['sample_id']}` | {fmt_n(int(r['n_cells']))} | "
                f"`{r['source_accession']}` | {r['preprocess_status']} | "
                f"{r['mapping_target']} |\n")

    f.write("\n## Notes\n\n")
    for topic, note in NOTES:
        f.write(f"- **{topic}.** {note}\n")

    f.write("\n## Summary statistics\n\n")
    for s in stats:
        f.write(f"- {s}\n")

# ------------------------------------------------------------
# Data sources (auto-appended to CSV / XLSX / DOCX / MD)
# ------------------------------------------------------------
DATA_SOURCES_ROWS = [
    ('PanSci', 'Zhang et al. Nature 2024', 'GSE247719', 'https://www.ncbi.nlm.nih.gov/geo/query/acc.cgi?acc=GSE247719', '2024-06', 'data/pansci/ (gitignored, see data_manifest/README.md §2.1)', 'sci-RNA-seq3 mouse pan-organ aging atlas'),
    ('EasySci brain', 'Sziraki et al. Nature Aging 2023', 'GSE212606', 'https://www.ncbi.nlm.nih.gov/geo/query/acc.cgi?acc=GSE212606', '2023-08', 'data/pansci/brain/ (gitignored)', 'EasySci, brain-optimised'),
    ('TMS (Tabula Muris Senis)', 'Schaum et al. Nature 2020', 'figshare:8273102', 'https://figshare.com/articles/dataset/8273102', '2020-07', 'data/tms/per_organ/ (gitignored)', 'Smart-seq2 FACS + 10x Chromium Droplet, 23 organs × 6 ages'),
    ('MCA2 (Mouse Cell Atlas 2.0)', 'Han et al. Cell 2022', 'GSE153562', 'https://www.ncbi.nlm.nih.gov/geo/query/acc.cgi?acc=GSE153562', '2022-09', 'data/mca2/ (gitignored)', 'Microwell-seq, fills aged-organ gaps'),
    ('Hammond microglia', 'Olah et al. Nat Commun 2020', 'GSE179358', 'https://www.ncbi.nlm.nih.gov/geo/query/acc.cgi?acc=GSE179358', '2020-04', 'data/microglia_aging/ (gitignored)', '10x microglia FACS-sorted'),
    ('Whole-body Array-seq mouse', 'Clevenger et al. Cell 2026', 'GSE248904', 'https://www.ncbi.nlm.nih.gov/geo/query/acc.cgi?acc=GSE248904', '2026-02', 'data/st/ + data/st/per_organ/ (gitignored)', 'Spatial reference for CMap projection'),
    ('CMap (single-cell → spatial mapping)', 'scripts/preprocess_new_sc.py + scripts/process_remaining_gaps.py', '-', '-', '-', 'models/cmap_output/{organ}/{sample}.h5ad (gitignored)', 'Projects each source single-cell sample onto Clevenger 2026 spatial reference'),
    ('gsMap', 'Liu et al. Nature Methods 2024', 'GitHub: JianYang-Lab/gsMap', 'https://github.com/JianYang-Lab/gsMap', 'v1.x', 'models/gsmap_age_output/{sample}/ (gitignored)', 'Spatial S-LDSC per (sample × annotation) Cauchy combination'),
    ('Derived provenance CSV', 'This work', '-', '-', '-', 'results/age_sample_provenance_per_source.csv', 'Rebuilt by scripts/build_age_provenance_per_source.py'),
]
DATA_SOURCES_COLS = ['data_source','citation','accession','url','version_snapshot','derived_path_in_repo','notes']
df_sources = pd.DataFrame(DATA_SOURCES_ROWS, columns=DATA_SOURCES_COLS)
df_sources.to_csv(OUT / 'supplement_table_mouse_1_data_sources.csv', index=False)

# Add as new sheet in xlsx
from openpyxl import load_workbook as _lw
from openpyxl.styles import Font as _F, PatternFill as _P, Alignment as _A
wb_ = _lw(str(OUT / 'supplement_table_mouse_1.xlsx'))
if 'Data Sources' in wb_.sheetnames:
    del wb_['Data Sources']
ws_ = wb_.create_sheet('Data Sources')
for c_idx, col in enumerate(df_sources.columns, 1):
    ws_.cell(row=1, column=c_idx, value=col)
for r_idx, row in enumerate(df_sources.itertuples(index=False), 2):
    for c_idx, v in enumerate(row, 1):
        ws_.cell(row=r_idx, column=c_idx, value=v)
for cell in ws_[1]:
    cell.fill = _P('solid', fgColor='1F4E79'); cell.font = _F(bold=True, color='FFFFFF', size=10)
    cell.alignment = _A(wrap_text=True, vertical='center')
ws_.freeze_panes = 'A2'
for cc in ws_.columns:
    m = max((len(str(c.value)) if c.value else 0) for c in cc)
    ws_.column_dimensions[cc[0].column_letter].width = min(max(m + 2, 12), 60)
for row in ws_.iter_rows(min_row=2):
    for cell in row:
        cell.font = _F(size=10); cell.alignment = _A(wrap_text=True, vertical='top')
wb_.save(str(OUT / 'supplement_table_mouse_1.xlsx'))

# Append section to docx
from docx import Document as _D
_doc = _D(str(OUT / 'supplement_table_mouse_1.docx'))
_doc.add_heading('Data sources', level=2)
add_df_table(_doc, df_sources, cap_widths=[2.6, 3.0, 2.0, 4.0, 2.0, 3.5, 5.0])
_doc.save(str(OUT / 'supplement_table_mouse_1.docx'))

# Append section to md
_md_p = OUT / 'supplement_table_mouse_1.md'
_md = _md_p.read_text()
if '## Data sources' not in _md:
    _md += '\n## Data sources\n\n'
    _md += '| Data source | Citation | Accession | URL | Version / snapshot | Path in repo | Notes |\n'
    _md += '|---|---|---|---|---|---|---|\n'
    for _, r in df_sources.iterrows():
        _md += (f"| **{r['data_source']}** | {r['citation']} | "
                f"`{r['accession']}` | {r['url']} | {r['version_snapshot']} | "
                f"`{r['derived_path_in_repo']}` | {r['notes']} |\n")
    _md_p.write_text(_md)

print(f"\nWrote:")
for p in sorted(OUT.glob("supplement_table_mouse_1*")):
    print(f"  {p.relative_to(REPO)} ({p.stat().st_size/1024:.1f} KB)")
