#!/usr/bin/env python
"""Build Supplementary Table Mouse 4 — final per-gene evidence matrix
combining GTEx eQTL + AD-SNP overlap + coloc-ABF + SharePro + SMR + HEIDI,
with explicit missing-analysis flags (Para 5 S6).

Outputs (5-format family):
  supplement_table_mouse_4.md
  supplement_table_mouse_4_evidence_matrix.csv    — Section 1: per-gene verdict
  supplement_table_mouse_4_per_method_raw.csv     — Section 2: raw per-method values
  supplement_table_mouse_4_notes.csv
  supplement_table_mouse_4_data_sources.csv       — sources
  supplement_table_mouse_4.xlsx                   — 4 sheets
  supplement_table_mouse_4.docx                   — landscape, formatted
"""
import os
from pathlib import Path
import numpy as np
import pandas as pd

REPO = Path(__file__).resolve().parents[1]
OUT  = REPO / "tables"
RES  = REPO / "results"

# Focal AD candidate genes from paragraph 5 (APP-pathway + classical hits)
FOCAL_GENES = ['CR1','PICALM','CD2AP','ADAM10','APP','SORL1','PSEN1','BIN1']

# Thresholds (NG convention)
THR_PP_H4    = 0.80          # strong single-causal coloc-ABF
THR_SHARE    = 0.50          # SharePro shared-causal threshold
THR_SMR_P    = 0.05 / 12     # Bonferroni on 12 SMR tests
THR_HEIDI_P  = 0.05          # HEIDI passes if > 0.05

# ---- Load all evidence streams ----
gtex     = pd.read_csv(RES / "gtex_eqtl_query_results.csv")
coloc    = pd.read_csv(RES / "coloc_FULL_eqtl_results.csv")
smr      = pd.read_csv(RES / "smr_colon_AD_results.csv")

# SharePro — per-gene file may or may not exist (CD2AP intentionally missing)
def share_for(gene):
    fp = RES / "sharepro" / f"{gene}_v3_result.sharepro.txt"
    if not fp.exists(): return None
    rows = []
    for li, line in enumerate(open(fp)):
        if li == 0: continue
        parts = line.strip().split('\t')
        if len(parts) >= 2:
            rows.append(float(parts[1]))
    return max(rows) if rows else None

# ---- Build per-gene evidence rows ----
ev_rows = []
for g in FOCAL_GENES:
    # GTEx colon
    cg = gtex[(gtex['gene']==g) & (gtex['tissue'].str.startswith('Colon'))]
    n_eqtl = int(cg['n_eqtl'].sum())
    best_p_eqtl = float(cg['best_p'].min()) if cg['best_p'].notna().any() else None
    ad_snp_is_colon_eqtl = bool(cg['ad_snp_is_eqtl'].any())

    # coloc-ABF (max PP.H4 across colon tissues)
    cc = coloc[(coloc['gene']==g) & (coloc['tissue'].str.startswith('Colon'))]
    best_pp_h4 = float(cc['PP_H4'].max()) if len(cc) else None
    best_pp_h3 = float(cc['PP_H3'].max()) if len(cc) else None

    # SharePro
    share = share_for(g)

    # SMR + HEIDI (best p_smr; corresponding heidi)
    ss = smr[smr['gene']==g]
    if len(ss):
        best_smr_p   = float(ss['p_smr'].min())
        row_at_best  = ss.loc[ss['p_smr'].idxmin()]
        best_heidi_p = float(row_at_best['p_heidi'])
    else:
        best_smr_p = best_heidi_p = None

    # Verdicts
    pass_coloc  = (best_pp_h4 is not None) and (best_pp_h4 >= THR_PP_H4)
    pass_share  = (share      is not None) and (share      >= THR_SHARE)
    pass_smr    = (best_smr_p is not None) and (best_smr_p <  THR_SMR_P)
    pass_heidi  = (best_heidi_p is not None) and (best_heidi_p > THR_HEIDI_P)

    # Missing-analysis flag
    missing = []
    if share is None:               missing.append("SharePro")
    if best_pp_h4 is None:          missing.append("coloc-ABF")
    if best_smr_p is None:          missing.append("SMR/HEIDI")

    # Final verdict
    if pass_coloc and pass_smr and pass_heidi:
        verdict = "✓✓ gold-standard causal (coloc + SMR + HEIDI pass)"
    elif pass_coloc or pass_share:
        if pass_coloc and pass_share:
            verdict = "✓✓ shared-causal (coloc + SharePro)"
        elif pass_coloc:
            verdict = "✓ shared-causal (single-causal coloc)"
        else:
            verdict = "✓ shared-causal (multi-causal SharePro rescue)"
    elif n_eqtl > 0 and ad_snp_is_colon_eqtl and "SharePro" in missing:
        verdict = (f"REQUIRES FOLLOW-UP: AD GWAS SNP overlaps a colon eQTL "
                   f"but multi-causal SharePro not yet run "
                   f"({' + '.join(missing)} missing)")
    elif n_eqtl > 0 and ad_snp_is_colon_eqtl:
        verdict = "AD SNP IS colon eQTL, but no coloc/SharePro support yet"
    elif n_eqtl > 0 and missing:
        verdict = f"REQUIRES FOLLOW-UP: has colon eQTL but missing {' + '.join(missing)}"
    elif n_eqtl > 0:
        verdict = "colon eQTL present but no coloc / SharePro / SMR support"
    else:
        verdict = "no colon eQTL"

    ev_rows.append({
        "gene": g,
        "GTEx_colon_n_eQTL":        n_eqtl,
        "GTEx_colon_best_p":        (None if best_p_eqtl is None
                                     else f"{best_p_eqtl:.2e}"),
        "AD_SNP_is_colon_eQTL":     "Yes" if ad_snp_is_colon_eqtl else "No",
        "coloc_ABF_best_PP_H4":     (None if best_pp_h4  is None
                                     else round(best_pp_h4, 4)),
        "coloc_ABF_best_PP_H3":     (None if best_pp_h3  is None
                                     else round(best_pp_h3, 4)),
        "SharePro_share":           (None if share is None
                                     else round(share, 4)),
        "SMR_best_p":               (None if best_smr_p is None
                                     else f"{best_smr_p:.2e}"),
        "HEIDI_p_at_best_SMR":      (None if best_heidi_p is None
                                     else f"{best_heidi_p:.2e}"),
        "passes_coloc_ABF":         "✓" if pass_coloc else "—",
        "passes_SharePro":          "✓" if pass_share else "—",
        "passes_SMR":               "✓" if pass_smr   else "—",
        "passes_HEIDI":             "✓" if pass_heidi else "—",
        "missing_analyses":         ", ".join(missing) if missing else "(complete)",
        "verdict":                  verdict,
    })

evidence = pd.DataFrame(ev_rows)

# ---- Section 2: per-method raw values (long format) ----
# Each row: gene × method × metric_name × value
raw_rows = []
for _, r in evidence.iterrows():
    g = r['gene']
    raw_rows.append({'gene': g, 'method': 'GTEx',     'metric': 'colon_n_eQTL',
                     'value': r['GTEx_colon_n_eQTL'], 'tissue': 'Colon (Trans+Sigmoid)'})
    raw_rows.append({'gene': g, 'method': 'GTEx',     'metric': 'colon_best_eQTL_p',
                     'value': r['GTEx_colon_best_p'], 'tissue': 'Colon (Trans+Sigmoid)'})
    raw_rows.append({'gene': g, 'method': 'GTEx',     'metric': 'AD_SNP_is_colon_eQTL',
                     'value': r['AD_SNP_is_colon_eQTL'], 'tissue': 'Colon (Trans+Sigmoid)'})
    raw_rows.append({'gene': g, 'method': 'coloc-ABF','metric': 'best_PP_H4',
                     'value': r['coloc_ABF_best_PP_H4'], 'tissue': 'Colon (Trans+Sigmoid)'})
    raw_rows.append({'gene': g, 'method': 'coloc-ABF','metric': 'best_PP_H3',
                     'value': r['coloc_ABF_best_PP_H3'], 'tissue': 'Colon (Trans+Sigmoid)'})
    raw_rows.append({'gene': g, 'method': 'SharePro', 'metric': 'share_probability',
                     'value': r['SharePro_share'],    'tissue': 'Colon (multi-causal)'})
    raw_rows.append({'gene': g, 'method': 'SMR',      'metric': 'best_SMR_p',
                     'value': r['SMR_best_p'],        'tissue': 'Colon (Trans+Sigmoid)'})
    raw_rows.append({'gene': g, 'method': 'HEIDI',    'metric': 'p_at_best_SMR',
                     'value': r['HEIDI_p_at_best_SMR'],'tissue': 'Colon (Trans+Sigmoid)'})
per_method = pd.DataFrame(raw_rows)

NOTES = [
    ("Thresholds",
     f"coloc-ABF PP.H4 ≥ {THR_PP_H4} = strong shared-causal; "
     f"SharePro share ≥ {THR_SHARE} = multi-causal shared signal; "
     f"SMR p < {THR_SMR_P:.4f} (Bonferroni on 12 tests); "
     f"HEIDI p > {THR_HEIDI_P} = pass (single-variant model consistent)."),
    ("Para 5 S5 verification",
     "PICALM SharePro share = 0.9977 (matches 0.998 in paragraph). "
     "CR1 SharePro share = 0.9991 (matches 0.999 in paragraph). "
     "ADAM10 SharePro share = 0.0005 → multi-causal NOT supported (matches paragraph)."),
    ("Para 5 S6 verification",
     "CD2AP has GTEx colon n_eQTL = 200, AD_SNP_is_colon_eQTL = Yes, "
     "but SharePro analysis was not run for CD2AP in this work (file "
     "CD2AP_v3_result.sharepro.txt does not exist). "
     "Verdict: 'REQUIRES FOLLOW-UP'. This is the gene called out by the "
     "paragraph as needing additional multi-causal colocalization work."),
    ("Why ADAM10 fails SharePro despite high SMR",
     "ADAM10 has SMR p = 4.0e-04 (Bonferroni-significant) but HEIDI p ≈ 0 "
     "(rejects single-variant model). SharePro share = 0.0005 (no shared "
     "causal under multi-causal model). Interpretation: ADAM10 has eQTL "
     "+ GWAS signal but the variants are NOT shared; SMR captures a "
     "linkage artifact that HEIDI correctly flags."),
    ("Final causal-candidate ranking (Colon)",
     "GOLD (coloc + SMR + HEIDI all pass): none. "
     "SILVER (coloc OR SharePro shared-causal): CR1 (both methods), "
     "PICALM (SharePro rescue). "
     "INCOMPLETE: CD2AP (SharePro not yet run). "
     "DROP: ADAM10 (multi-causal contradicts shared-causal interpretation). "
     "NO COLON EQTL: SORL1, PSEN1 (zero significant colon variants in GTEx)."),
    ("How to extend",
     "To run SharePro for CD2AP: place GWAS + eQTL summary statistics + "
     "LD matrix in results/sharepro/CD2AP_v3_{gwas,eqtl,ld}.txt, then run "
     "the SharePro CLI as in `notebooks/21_eqtl_coloc_analysis.ipynb`. "
     "Output will be results/sharepro/CD2AP_v3_result.sharepro.txt and "
     "this builder will pick it up automatically on re-run."),
]
df_notes = pd.DataFrame(NOTES, columns=["topic","note"])

# ---- Write CSVs ----
evidence.to_csv  (OUT / "supplement_table_mouse_4_evidence_matrix.csv", index=False)
per_method.to_csv(OUT / "supplement_table_mouse_4_per_method_raw.csv",   index=False)
df_notes.to_csv  (OUT / "supplement_table_mouse_4_notes.csv",            index=False)

# ---- XLSX ----
xlsx_path = OUT / "supplement_table_mouse_4.xlsx"
with pd.ExcelWriter(xlsx_path, engine="openpyxl") as xw:
    evidence.to_excel  (xw, sheet_name="1 - Evidence matrix",      index=False)
    per_method.to_excel(xw, sheet_name="2 - Per-method raw values", index=False)
    df_notes.to_excel  (xw, sheet_name="Notes",                    index=False)

from openpyxl import load_workbook
from openpyxl.styles import Font, PatternFill, Alignment
wb = load_workbook(xlsx_path)
HF, HFNT = PatternFill("solid", fgColor="1F4E79"), Font(bold=True, color="FFFFFF", size=10)
BF, AL   = Font(size=10), Alignment(wrap_text=True, vertical="top")
for sheet in wb.sheetnames:
    ws = wb[sheet]
    for cell in ws[1]:
        cell.fill = HF; cell.font = HFNT
        cell.alignment = Alignment(wrap_text=True, vertical="center")
    ws.freeze_panes = "A2"
    for cc in ws.columns:
        m = max((len(str(c.value)) if c.value else 0) for c in cc)
        ws.column_dimensions[cc[0].column_letter].width = min(max(m + 2, 12), 55)
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.font = BF; cell.alignment = AL
wb.save(xlsx_path)

# ---- DOCX ----
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
def shade_cell(cell, hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto"); shd.set(qn("w:fill"), hex); tcPr.append(shd)
def add_df_table(doc, df_in, hex="1F4E79", cap_widths=None):
    table = doc.add_table(rows=1+len(df_in), cols=len(df_in.columns))
    table.style = "Light Grid Accent 1"
    for i, col in enumerate(df_in.columns):
        c = table.rows[0].cells[i]; c.text = str(col)
        for run in c.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); run.font.size = Pt(9)
        shade_cell(c, hex)
    for r, (_, row) in enumerate(df_in.iterrows(), 1):
        for i, v in enumerate(row.values):
            table.rows[r].cells[i].text = "" if pd.isna(v) else str(v)
            for run in table.rows[r].cells[i].paragraphs[0].runs:
                run.font.size = Pt(8.5)
    if cap_widths:
        for i, w in enumerate(cap_widths):
            if w is None: continue
            for cell in table.columns[i].cells:
                cell.width = Cm(w)

doc = Document()
section = doc.sections[0]; section.orientation = WD_ORIENT.LANDSCAPE
section.page_width, section.page_height = section.page_height, section.page_width
section.top_margin = section.bottom_margin = Cm(1.2)
section.left_margin = section.right_margin = Cm(1.2)
for name in ("Normal","Heading 1","Heading 2","Heading 3"):
    if name in doc.styles: doc.styles[name].font.name = "Arial"

doc.add_heading("Supplementary Table Mouse 4", level=1)
doc.add_paragraph(
    "Final per-gene evidence matrix combining GTEx colon eQTL, AD-SNP × eQTL "
    "overlap, single-causal coloc-ABF (full summary statistics), multi-causal "
    "SharePro, SMR + HEIDI, with explicit missing-analysis flags (Para 5 S5–S6). "
    "Highlights the CD2AP follow-up gap called out in S6.").runs[0].font.size = Pt(10)

doc.add_heading("1 | Final evidence matrix (8 candidate genes)", level=2)
add_df_table(doc, evidence)

doc.add_heading("2 | Per-method raw values (long format)", level=2)
add_df_table(doc, per_method)

doc.add_heading("Notes", level=2)
for topic, note in NOTES:
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(f"{topic}. "); r1.font.bold = True; r1.font.size = Pt(10)
    r2 = p.add_run(note);          r2.font.size = Pt(10)

doc.save(OUT / "supplement_table_mouse_4.docx")

# ---- Markdown ----
with open(OUT / "supplement_table_mouse_4.md", "w") as f:
    f.write("## Supplementary Table Mouse 4 | Final per-gene evidence matrix "
            "(Para 5 S5–S6)\n\n")
    f.write("Combined evidence from **GTEx eQTL + AD-SNP overlap + coloc-ABF "
            "(full summary stats) + SharePro + SMR + HEIDI** for the 8 focal "
            "AD risk genes nominated by the colon spatial analysis. "
            "Each method has its own threshold (NG convention); the **verdict** "
            "column gives the integrated call, and **missing_analyses** is the "
            "explicit gap that the paragraph S6 calls out for CD2AP.\n\n")

    f.write("### 1 | Evidence matrix\n\n")
    f.write("| Gene | colon n_eQTL | colon best_p | AD-SNP=colon eQTL | "
            "coloc PP.H4 | coloc PP.H3 | SharePro share | SMR best p | HEIDI p | "
            "coloc✓ | SharePro✓ | SMR✓ | HEIDI✓ | Missing | Verdict |\n")
    f.write("|---|---:|---:|---|---:|---:|---:|---:|---:|:---:|:---:|:---:|:---:|---|---|\n")
    for _, r in evidence.iterrows():
        def cell(v):
            return "" if v is None else str(v)
        f.write(f"| **{r['gene']}** | {r['GTEx_colon_n_eQTL']} | "
                f"{cell(r['GTEx_colon_best_p'])} | {r['AD_SNP_is_colon_eQTL']} | "
                f"{cell(r['coloc_ABF_best_PP_H4'])} | {cell(r['coloc_ABF_best_PP_H3'])} | "
                f"{cell(r['SharePro_share'])} | {cell(r['SMR_best_p'])} | "
                f"{cell(r['HEIDI_p_at_best_SMR'])} | "
                f"{r['passes_coloc_ABF']} | {r['passes_SharePro']} | "
                f"{r['passes_SMR']} | {r['passes_HEIDI']} | "
                f"{r['missing_analyses']} | {r['verdict']} |\n")

    f.write("\n### 2 | Per-method raw values (long format)\n\n")
    f.write("| Gene | Method | Metric | Value | Tissue |\n")
    f.write("|---|---|---|---:|---|\n")
    for _, r in per_method.iterrows():
        f.write(f"| {r['gene']} | {r['method']} | {r['metric']} | "
                f"{'' if r['value'] is None else r['value']} | {r['tissue']} |\n")

    f.write("\n## Notes\n\n")
    for topic, note in NOTES:
        f.write(f"- **{topic}.** {note}\n")

    f.write("\n## Summary statistics\n\n")
    n_complete = (evidence['missing_analyses'] == '(complete)').sum()
    n_followup = evidence['missing_analyses'].str.contains('SharePro').sum()
    f.write(f"- Focal genes: {len(evidence)}\n")
    f.write(f"- Complete pipeline (no missing analyses): {n_complete}\n")
    f.write(f"- Genes flagged for follow-up: {n_followup} "
            f"({', '.join(evidence.loc[evidence['missing_analyses'].str.contains('SharePro'),'gene'].tolist())})\n")
    f.write(f"- Genes passing coloc-ABF (PP.H4 ≥ {THR_PP_H4}): "
            f"{int((evidence['passes_coloc_ABF']=='✓').sum())} "
            f"({', '.join(evidence.loc[evidence['passes_coloc_ABF']=='✓','gene'].tolist())})\n")
    f.write(f"- Genes passing SharePro (share ≥ {THR_SHARE}): "
            f"{int((evidence['passes_SharePro']=='✓').sum())} "
            f"({', '.join(evidence.loc[evidence['passes_SharePro']=='✓','gene'].tolist())})\n")
    f.write(f"- Genes passing SMR (Bonferroni): "
            f"{int((evidence['passes_SMR']=='✓').sum())}\n")
    f.write(f"- Genes passing HEIDI: "
            f"{int((evidence['passes_HEIDI']=='✓').sum())}\n")

# ---- Data sources block (auto-appended) ----
DATA_SOURCES_ROWS = [
    ('GTEx Portal v8 eQTL query', 'GTEx Consortium Science 2020', 'v8',
     'https://gtexportal.org/home/', 'queried 2026-04 (snapshot)',
     'results/gtex_eqtl_query_results.csv',
     'Per (gene × tissue) significant cis-eQTL counts + best p + AD-SNP-is-eQTL flag'),
    ('Bellenguez 2022 AD GWAS', 'Bellenguez et al. Nat Genet 2022', 'GCST90027158',
     'https://www.ebi.ac.uk/gwas/studies/GCST90027158', '2022-04 (GRCh38)',
     'data/gwas/Bellenguez2022_AD_withN.tsv.gz (gitignored)',
     'AD GWAS lead SNPs per locus'),
    ('coloc-ABF (full summary stats)', 'Giambartolomei et al. PLoS Genet 2014', '-',
     'https://github.com/chr1swallace/coloc', 'coloc v5.x',
     'results/coloc_FULL_eqtl_results.csv',
     'Single-causal coloc using full eQTL summary stats; PP.H0–H4 per (gene × colon tissue)'),
    ('SharePro multi-causal coloc', 'Zhang et al. Genome Biol 2023', '-',
     'https://github.com/zhwm/SharePro_coloc', 'SharePro v5.0.0',
     'results/sharepro/{GENE}_v3_result.sharepro.txt',
     'Multi-causal shared-component coloc; produces "share" probability per credible set'),
    ('SMR + HEIDI', 'Zhu et al. Nat Genet 2016', '-',
     'https://yanglab.westlake.edu.cn/software/smr/', 'SMR v1.3.1',
     'results/smr_colon_AD_results.csv',
     'Summary-data-based Mendelian Randomization with HEIDI heterogeneity test'),
    ('Pipeline notebook', 'This work', '-', '-', '-',
     'notebooks/21_eqtl_coloc_analysis.ipynb (+ _executed)',
     'Orchestrates GTEx query + coloc-ABF + SharePro + SMR for the 8 focal genes'),
    ('Comparison data: sig-only vs full', 'Supplementary Fig (Para 5 S4)', '-',
     '-', '-', 'results/coloc_sig_vs_full_comparison.csv',
     'Sister CSV showing PP.H4 drop when switching from sig-only to full summary'),
]
DATA_SOURCES_COLS = ['data_source','citation','accession','url','version_snapshot',
                     'derived_path_in_repo','notes']
df_sources = pd.DataFrame(DATA_SOURCES_ROWS, columns=DATA_SOURCES_COLS)
df_sources.to_csv(OUT / 'supplement_table_mouse_4_data_sources.csv', index=False)

# add to xlsx
from openpyxl import load_workbook as _lw
wb_ = _lw(str(OUT / 'supplement_table_mouse_4.xlsx'))
if 'Data Sources' in wb_.sheetnames: del wb_['Data Sources']
ws_ = wb_.create_sheet('Data Sources')
for c_idx, col in enumerate(df_sources.columns, 1):
    ws_.cell(row=1, column=c_idx, value=col)
for r_idx, row in enumerate(df_sources.itertuples(index=False), 2):
    for c_idx, v in enumerate(row, 1):
        ws_.cell(row=r_idx, column=c_idx, value=v)
for cell in ws_[1]:
    cell.fill = PatternFill('solid', fgColor='1F4E79')
    cell.font = Font(bold=True, color='FFFFFF', size=10)
    cell.alignment = Alignment(wrap_text=True, vertical='center')
ws_.freeze_panes = 'A2'
for cc in ws_.columns:
    m = max((len(str(c.value)) if c.value else 0) for c in cc)
    ws_.column_dimensions[cc[0].column_letter].width = min(max(m + 2, 12), 60)
for row in ws_.iter_rows(min_row=2):
    for cell in row:
        cell.font = Font(size=10); cell.alignment = Alignment(wrap_text=True, vertical='top')
wb_.save(str(OUT / 'supplement_table_mouse_4.xlsx'))

# add to docx
doc2 = Document(str(OUT / 'supplement_table_mouse_4.docx'))
doc2.add_heading('Data sources', level=2)
add_df_table(doc2, df_sources, cap_widths=[2.6,3.0,2.0,4.0,2.0,3.5,5.0])
doc2.save(str(OUT / 'supplement_table_mouse_4.docx'))

# add to md
md_p = OUT / 'supplement_table_mouse_4.md'
md = md_p.read_text()
md += "\n## Data sources\n\n"
md += "| Data source | Citation | Accession | URL | Version / snapshot | Path in repo | Notes |\n"
md += "|---|---|---|---|---|---|---|\n"
for _, r in df_sources.iterrows():
    md += (f"| **{r['data_source']}** | {r['citation']} | "
           f"`{r['accession']}` | {r['url']} | {r['version_snapshot']} | "
           f"`{r['derived_path_in_repo']}` | {r['notes']} |\n")
md_p.write_text(md)

print("Wrote:")
for p in sorted(OUT.glob("supplement_table_mouse_4*")):
    print(f"  {p.relative_to(REPO)} ({p.stat().st_size/1024:.1f} KB)")
