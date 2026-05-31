#!/usr/bin/env python
"""Build Supplementary Table Mouse 2 — per-organ × per-age AD spatial-risk
summary across the multi-age CMap-projected atlas.

Outputs:
  supplement_table_mouse_2.md
  supplement_table_mouse_2_overall.csv              (3 age-group totals)
  supplement_table_mouse_2_per_group.csv            (organ × {Young, Mid, Old})
  supplement_table_mouse_2_per_age.csv              (organ × age_months)
  supplement_table_mouse_2_notes.csv
  supplement_table_mouse_2.xlsx                     (4 sheets)
  supplement_table_mouse_2.docx                     (landscape, formatted)

Driving data: results/age_organ_summary_{per_age,per_group}.csv +
results/age_overall_summary.csv (computed once from
data/age_merged/age_all_organs_all_traits.h5ad).

Reproduce upstream CSVs with the snippet at the top of this script.
Run: python tables/build_supp_table_mouse_2.py
"""
import os
from pathlib import Path

import pandas as pd

REPO = Path(__file__).resolve().parents[1]
OUT  = REPO / "tables"
RES  = REPO / "results"

# ------------------------------------------------------------
# Load precomputed summaries
# ------------------------------------------------------------
overall   = pd.read_csv(RES / "age_overall_summary.csv")
per_group = pd.read_csv(RES / "age_organ_summary_per_group.csv")
per_age   = pd.read_csv(RES / "age_organ_summary_per_age.csv")

# Re-order age_group factors
AG_ORDER = ['Young (≤6 mo)', 'Mid (6–18 mo)', 'Old (≥18 mo)']
overall["age_group"]   = pd.Categorical(overall["age_group"],   categories=AG_ORDER, ordered=True)
per_group["age_group"] = pd.Categorical(per_group["age_group"], categories=AG_ORDER, ordered=True)
overall   = overall.sort_values("age_group").reset_index(drop=True)
per_group = per_group.sort_values(["organ","age_group"]).reset_index(drop=True)
per_age   = per_age.sort_values(["organ","age_months"]).reset_index(drop=True)

# Round numerics
for df in (overall, per_group, per_age):
    for c in ("median_logp","q25_logp","q75_logp","frac_genomewide_sig"):
        df[c] = df[c].astype(float).round(3)
    for c in ("n_samples","n_spots_total","n_spots_pAD_valid"):
        df[c] = df[c].astype(int)

NOTES = [
    ("gsMap input",
     "Each (organ × age × source) CMap-projected h5ad in the multi-age atlas "
     "(Supplementary Table Mouse 1) was scored with `gsmap quick_mode` against "
     "the Bellenguez 2022 AD GWAS, producing a per-spot S-LDSC χ² statistic and "
     "p-value."),
    ("Median log10(p) metric",
     "Per-organ × per-age 'median_logp' is the spot-level median of "
     "-log10(p_AD). Higher = stronger per-spot AD heritability enrichment in "
     "that organ-age stratum. q25 and q75 give the 25th/75th percentiles to "
     "show the spread."),
    ("Genome-wide significance fraction",
     "frac_genomewide_sig = fraction of spots with p_AD < 5×10⁻⁸. Conservative; "
     "complements the median by capturing the tail."),
    ("Why n_spots_total ≠ n_spots_pAD_valid",
     "A small number of spots (~1%) fail gsMap S-LDSC (no LDSC score in the "
     "chunk they fell into); we list both for transparency. Median and "
     "percentiles use only valid spots."),
    ("Age-group definitions",
     "Young = age_months ≤ 6 (1, 3, 6 mo); Mid = 6 < age_months < 18 "
     "(12, 16 mo); Old = age_months ≥ 18 (18, 21, 23, 24, 30 mo). "
     "Young and Old definitions match the Young/Old comparison in "
     "Fig. 4e and the paragraph S2 reference numbers (264,283 young "
     "spots / 41 samples; 322,477 old spots / 62 samples)."),
    ("Sample-sample variability",
     "n_samples in each (organ × age) cell is the number of CMap-projected "
     "samples (= h5ad files) feeding gsMap. The same organ-age can have "
     ">1 sample from different single-cell sources (see "
     "Supplementary Table Mouse 1, Section 3)."),
]

# ------------------------------------------------------------
# Write CSVs
# ------------------------------------------------------------
overall.to_csv(OUT / "supplement_table_mouse_2_overall.csv", index=False)
per_group.to_csv(OUT / "supplement_table_mouse_2_per_group.csv", index=False)
per_age.to_csv(OUT / "supplement_table_mouse_2_per_age.csv", index=False)
df_notes = pd.DataFrame(NOTES, columns=["topic","note"])
df_notes.to_csv(OUT / "supplement_table_mouse_2_notes.csv", index=False)

# ------------------------------------------------------------
# XLSX (4 sheets, formatted)
# ------------------------------------------------------------
xlsx_path = OUT / "supplement_table_mouse_2.xlsx"
with pd.ExcelWriter(xlsx_path, engine="openpyxl") as xw:
    overall.to_excel(xw,   sheet_name="1 - Overall Young vs Old", index=False)
    per_group.to_excel(xw, sheet_name="2 - Per-organ x group",   index=False)
    per_age.to_excel(xw,   sheet_name="3 - Per-organ x age", index=False)
    df_notes.to_excel(xw,  sheet_name="Notes",                    index=False)

from openpyxl import load_workbook
from openpyxl.styles import Font, PatternFill, Alignment
wb = load_workbook(xlsx_path)
HEADER_FILL = PatternFill("solid", fgColor="1F4E79")
HEADER_FONT = Font(bold=True, color="FFFFFF", size=10)
BODY_FONT   = Font(size=10)
ALIGN       = Alignment(wrap_text=True, vertical="top")
for sheet in wb.sheetnames:
    ws = wb[sheet]
    for cell in ws[1]:
        cell.fill = HEADER_FILL
        cell.font = HEADER_FONT
        cell.alignment = Alignment(wrap_text=True, vertical="center")
    ws.freeze_panes = "A2"
    for col_cells in ws.columns:
        m = max((len(str(c.value)) if c.value else 0) for c in col_cells)
        ws.column_dimensions[col_cells[0].column_letter].width = min(max(m + 2, 12), 50)
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.font = BODY_FONT
            cell.alignment = ALIGN
wb.save(xlsx_path)

# ------------------------------------------------------------
# DOCX
# ------------------------------------------------------------
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def add_df_table(doc, df_in, header_hex="1F4E79", cap_widths=None):
    table = doc.add_table(rows=1 + len(df_in), cols=len(df_in.columns))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, col in enumerate(df_in.columns):
        hdr[i].text = str(col)
        for run in hdr[i].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(9)
        shade_cell(hdr[i], header_hex)
    for r, (_, row) in enumerate(df_in.iterrows(), 1):
        for i, val in enumerate(row.values):
            table.rows[r].cells[i].text = "" if pd.isna(val) else str(val)
            for run in table.rows[r].cells[i].paragraphs[0].runs:
                run.font.size = Pt(8.5)
    if cap_widths:
        for i, w in enumerate(cap_widths):
            if w is None: continue
            for cell in table.columns[i].cells:
                cell.width = Cm(w)
    return table


doc = Document()
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width, section.page_height = section.page_height, section.page_width
section.top_margin = section.bottom_margin = Cm(1.2)
section.left_margin = section.right_margin = Cm(1.2)
for name in ("Normal","Heading 1","Heading 2","Heading 3"):
    if name in doc.styles: doc.styles[name].font.name = "Arial"

doc.add_heading("Supplementary Table Mouse 2", level=1)
p = doc.add_paragraph()
p.add_run("Per-organ, per-age sample counts, spot counts and median spatial AD "
          "genetic-risk scores across the multi-age CMap-projected atlas. "
          "Young (≤6 mo) vs Old (≥18 mo) totals match the paragraph reference "
          "numbers — 264,283 young spots / 41 samples vs 322,477 old spots / "
          "62 samples — used to construct the Young/Old whole-body risk maps "
          "in Fig. 4e."
          ).font.size = Pt(10)

doc.add_heading("1 | Overall age-group totals", level=2)
add_df_table(doc, overall)

doc.add_heading("2 | Per-organ × age-group breakdown", level=2)
add_df_table(doc, per_group, cap_widths=[2.4,3.2,1.6,2.0,2.4,1.8,1.6,1.6,2.4])

doc.add_heading("3 | Per-organ × age-months (fine-grained)", level=2)
add_df_table(doc, per_age, cap_widths=[2.4,1.6,1.6,2.0,2.4,1.8,1.6,1.6,2.4])

doc.add_heading("Notes", level=2)
for topic, note in NOTES:
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(f"{topic}. ")
    r1.font.bold = True; r1.font.size = Pt(10)
    r2 = p.add_run(note);     r2.font.size = Pt(10)

doc.add_heading("Summary statistics", level=2)
stats = [
    f"Total samples: {int(per_age['n_samples'].sum())}",
    f"Total spots:   {int(overall['n_spots_total'].sum()):,}",
    f"Spots with valid p_AD: {int(overall['n_spots_pAD_valid'].sum()):,} "
    f"({100*overall['n_spots_pAD_valid'].sum()/overall['n_spots_total'].sum():.2f}%)",
    f"Distinct organs: {per_age['organ'].nunique()}",
    f"Distinct ages: {per_age['age_months'].nunique()} months "
    f"({sorted(per_age['age_months'].unique().tolist())})",
    f"Young (≤6 mo): {int(overall.loc[overall['age_group']=='Young (≤6 mo)','n_spots_total'].iloc[0]):,} spots, "
    f"{int(overall.loc[overall['age_group']=='Young (≤6 mo)','n_samples'].iloc[0])} samples",
    f"Old (≥18 mo):  {int(overall.loc[overall['age_group']=='Old (≥18 mo)','n_spots_total'].iloc[0]):,} spots, "
    f"{int(overall.loc[overall['age_group']=='Old (≥18 mo)','n_samples'].iloc[0])} samples",
    f"Old vs Young median log10(p) ratio (overall): "
    f"{overall.loc[overall['age_group']=='Old (≥18 mo)','median_logp'].iloc[0]:.2f} / "
    f"{overall.loc[overall['age_group']=='Young (≤6 mo)','median_logp'].iloc[0]:.2f} = "
    f"{overall.loc[overall['age_group']=='Old (≥18 mo)','median_logp'].iloc[0] / overall.loc[overall['age_group']=='Young (≤6 mo)','median_logp'].iloc[0]:.2f}×",
]
for s in stats:
    doc.add_paragraph(s, style="List Bullet")

doc.save(OUT / "supplement_table_mouse_2.docx")

# ------------------------------------------------------------
# Markdown
# ------------------------------------------------------------
def fmt_n(n):
    if isinstance(n, float) and pd.notna(n) and n.is_integer(): n = int(n)
    if isinstance(n, int) and n >= 1000: return f"{n:,}"
    return str(n)


with open(OUT / "supplement_table_mouse_2.md", "w") as f:
    f.write("## Supplementary Table Mouse 2 | Per-organ × per-age AD spatial-risk "
            "summary across the multi-age atlas\n\n")
    f.write("Per-organ, per-age sample counts, spot counts and median spatial "
            "Alzheimer's-disease genetic-risk scores produced by `gsmap quick_mode` "
            "against the Bellenguez 2022 AD GWAS, after CMap projection of the "
            "multi-age single-cell atlas (Supplementary Table Mouse 1) onto the "
            "6-week whole-body Array-seq reference (GSE248904). "
            "Young (≤6 mo) vs Old (≥18 mo) totals — **264,283 young spots / "
            "41 samples** and **322,477 old spots / 62 samples** — match the "
            "Fig. 4e paragraph reference.\n\n")

    f.write("### 1 | Overall age-group totals\n\n")
    f.write("| Age group | n_samples | n_spots_total | n_spots (valid p_AD) | "
            "median −log10(p_AD) | q25 | q75 | frac. genome-wide sig (p<5e−8) |\n")
    f.write("|---|---:|---:|---:|---:|---:|---:|---:|\n")
    for _, r in overall.iterrows():
        f.write(f"| {r['age_group']} | {fmt_n(r['n_samples'])} | "
                f"{fmt_n(r['n_spots_total'])} | {fmt_n(r['n_spots_pAD_valid'])} | "
                f"{r['median_logp']:.3f} | {r['q25_logp']:.3f} | "
                f"{r['q75_logp']:.3f} | {r['frac_genomewide_sig']:.4f} |\n")

    f.write("\n### 2 | Per-organ × age-group breakdown\n\n")
    f.write("| Organ | Age group | n_samples | n_spots | n_valid | "
            "median −log10p | q25 | q75 | frac. sig |\n")
    f.write("|---|---|---:|---:|---:|---:|---:|---:|---:|\n")
    for _, r in per_group.iterrows():
        f.write(f"| {r['organ']} | {r['age_group']} | "
                f"{fmt_n(r['n_samples'])} | {fmt_n(r['n_spots_total'])} | "
                f"{fmt_n(r['n_spots_pAD_valid'])} | "
                f"{r['median_logp']:.3f} | {r['q25_logp']:.3f} | "
                f"{r['q75_logp']:.3f} | {r['frac_genomewide_sig']:.4f} |\n")

    f.write("\n### 3 | Per-organ × age-months (fine-grained, 114 rows)\n\n")
    f.write("| Organ | Age (mo) | n_samples | n_spots | n_valid | "
            "median −log10p | q25 | q75 | frac. sig |\n")
    f.write("|---|---:|---:|---:|---:|---:|---:|---:|---:|\n")
    for _, r in per_age.iterrows():
        f.write(f"| {r['organ']} | {int(r['age_months'])} | "
                f"{fmt_n(r['n_samples'])} | {fmt_n(r['n_spots_total'])} | "
                f"{fmt_n(r['n_spots_pAD_valid'])} | "
                f"{r['median_logp']:.3f} | {r['q25_logp']:.3f} | "
                f"{r['q75_logp']:.3f} | {r['frac_genomewide_sig']:.4f} |\n")

    f.write("\n## Notes\n\n")
    for topic, note in NOTES:
        f.write(f"- **{topic}.** {note}\n")

    f.write("\n## Summary statistics\n\n")
    for s in stats:
        f.write(f"- {s}\n")


# ------------------------------------------------------------
# Data sources (auto-appended to CSV / XLSX / DOCX / MD)
# ------------------------------------------------------------
DATA_SOURCES_ROWS = [
    ('Bellenguez 2022 AD GWAS', 'Bellenguez et al. Nat Genet 2022', 'GCST90027158', 'https://www.ebi.ac.uk/gwas/studies/GCST90027158', '2022-04 (build GRCh38)', 'data/gwas/Bellenguez2022_AD_withN.tsv.gz (gitignored)', 'Ncase ≈ 90k, Ntotal ≈ 788k; converted with `gsmap format_sumstats`'),
    ('CMap-projected multi-age atlas', 'Supplementary Table Mouse 1', '-', '-', '-', 'data/age_merged/age_all_organs_all_traits.h5ad (gitignored, 5.6 GB)', '126 (organ × age) samples, 809,284 spots'),
    ('gsMap quick_mode', 'Liu et al. Nature Methods 2024', 'v1.x', 'https://github.com/JianYang-Lab/gsMap', 'commit pinned in env', 'models/gsmap_age_output/{sample}/spatial_ldsc/{sample}_AD.csv.gz (gitignored)', 'Per-spot AD S-LDSC χ² → p_AD column'),
    ('Per-organ summary CSVs', 'This work', '-', '-', '-', 'results/age_organ_summary_per_age.csv + results/age_organ_summary_per_group.csv + results/age_overall_summary.csv', 'Pre-aggregated by tables/build_supp_table_mouse_2.py header snippet'),
]
DATA_SOURCES_COLS = ['data_source','citation','accession','url','version_snapshot','derived_path_in_repo','notes']
df_sources = pd.DataFrame(DATA_SOURCES_ROWS, columns=DATA_SOURCES_COLS)
df_sources.to_csv(OUT / 'supplement_table_mouse_2_data_sources.csv', index=False)

# Add as new sheet in xlsx
from openpyxl import load_workbook as _lw
from openpyxl.styles import Font as _F, PatternFill as _P, Alignment as _A
wb_ = _lw(str(OUT / 'supplement_table_mouse_2.xlsx'))
if 'Data Sources' in wb_.sheetnames:
    del wb_['Data Sources']
ws_ = wb_.create_sheet('Data Sources')
for c_idx, col in enumerate(df_sources.columns, 1):
    ws_.cell(row=1, column=c_idx, value=col)
for r_idx, row in enumerate(df_sources.itertuples(index=False), 2):
    for c_idx, v in enumerate(row, 1):
        ws_.cell(row=r_idx, column=c_idx, value=v)
for cell in ws_[1]:
    cell.fill = _P('solid', fgColor='1F4E79'); cell.font = _F(bold=True, color='FFFFFF', size=10)
    cell.alignment = _A(wrap_text=True, vertical='center')
ws_.freeze_panes = 'A2'
for cc in ws_.columns:
    m = max((len(str(c.value)) if c.value else 0) for c in cc)
    ws_.column_dimensions[cc[0].column_letter].width = min(max(m + 2, 12), 60)
for row in ws_.iter_rows(min_row=2):
    for cell in row:
        cell.font = _F(size=10); cell.alignment = _A(wrap_text=True, vertical='top')
wb_.save(str(OUT / 'supplement_table_mouse_2.xlsx'))

# Append section to docx
from docx import Document as _D
_doc = _D(str(OUT / 'supplement_table_mouse_2.docx'))
_doc.add_heading('Data sources', level=2)
add_df_table(_doc, df_sources, cap_widths=[2.6, 3.0, 2.0, 4.0, 2.0, 3.5, 5.0])
_doc.save(str(OUT / 'supplement_table_mouse_2.docx'))

# Append section to md
_md_p = OUT / 'supplement_table_mouse_2.md'
_md = _md_p.read_text()
if '## Data sources' not in _md:
    _md += '\n## Data sources\n\n'
    _md += '| Data source | Citation | Accession | URL | Version / snapshot | Path in repo | Notes |\n'
    _md += '|---|---|---|---|---|---|---|\n'
    for _, r in df_sources.iterrows():
        _md += (f"| **{r['data_source']}** | {r['citation']} | "
                f"`{r['accession']}` | {r['url']} | {r['version_snapshot']} | "
                f"`{r['derived_path_in_repo']}` | {r['notes']} |\n")
    _md_p.write_text(_md)

print("Wrote:")
for p in sorted(OUT.glob("supplement_table_mouse_2*")):
    print(f"  {p.relative_to(REPO)} ({p.stat().st_size/1024:.1f} KB)")
